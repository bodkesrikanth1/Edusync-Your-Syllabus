import io
import os
import functools
from flask import Flask, render_template, request, redirect, url_for, flash, session, g
from config import Config
from db import insert_syllabus, insert_unit, insert_topic, insert_video, fetch_units_topics_videos
from db import create_user, get_user_by_email, get_user_by_id, get_conn
from nlp import extract_topics_per_unit
from youtube_api import search_and_rank
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from docx import Document as DocxDocument

ALLOWED_EXT = {"txt", "docx", "doc"}  # .doc will be rejected gently (see below)
MAX_UPLOAD_BYTES = 2 * 1024 * 1024  # 2 MB max per file (adjust as needed)

app = Flask(__name__)
app.config.from_object(Config)
app.secret_key = Config.SECRET_KEY

# Error handlers
@app.errorhandler(500)
def internal_error(error):
    import traceback
    print(f"500 Internal Server Error: {error}")
    print(traceback.format_exc())
    return render_template('error.html', error="Internal Server Error", code=500), 500

@app.errorhandler(404)
def not_found(error):
    return render_template('error.html', error="Page Not Found", code=404), 404

@app.errorhandler(Exception)
def handle_exception(error):
    import traceback
    print(f"Unhandled exception: {error}")
    print(traceback.format_exc())
    return render_template('error.html', error="Application Error", code=500), 500

# ---------- Helpers ----------
def login_required(view):
    @functools.wraps(view)
    def wrapped_view(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login', next=request.path))
        return view(*args, **kwargs)
    return wrapped_view

@app.before_request
def load_logged_in_user():
    """Load authenticated user into g object for every request."""
    try:
        user_id = session.get('user_id')
        if user_id is None:
            g.user = None
        else:
            try:
                g.user = get_user_by_id(user_id)
                if g.user is None:
                    # User was deleted from database
                    session.clear()
            except Exception as e:
                print(f"Error loading user {user_id}: {e}")
                g.user = None
                session.clear()
    except Exception as e:
        print(f"before_request error: {e}")
        g.user = None

def allowed_file(filename: str):
    if not filename:
        return False
    ext = filename.rsplit('.', 1)[-1].lower()
    return ext in ALLOWED_EXT

def extract_text_from_docx_bytes(b: bytes) -> str:
    """
    Extract text from .docx binary content using python-docx.
    """
    f = io.BytesIO(b)
    doc = DocxDocument(f)
    paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    # also try tables if present
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text and cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs).strip()

def extract_text_from_upload(file_storage) -> (str, str):
    """
    Accepts a werkzeug FileStorage. Returns tuple (text, detected_type)
    detected_type: 'txt'|'docx'|'doc'|'unknown'
    """
    filename = secure_filename(file_storage.filename or "")
    if not filename:
        return "", "none"
    ext = filename.rsplit('.', 1)[-1].lower()
    # limit size
    data = file_storage.stream.read(MAX_UPLOAD_BYTES + 1)
    if len(data) > MAX_UPLOAD_BYTES:
        raise ValueError("File too large. Max 2 MB.")
    if ext == "txt":
        try:
            # try decode as utf-8 with fallback to latin-1
            try:
                text = data.decode("utf-8")
            except UnicodeDecodeError:
                text = data.decode("latin-1")
            return text.strip(), "txt"
        except Exception:
            return "", "txt"
    elif ext == "docx":
        try:
            text = extract_text_from_docx_bytes(data)
            return text, "docx"
        except Exception as e:
            print("docx parse error:", e)
            return "", "docx"
    elif ext == "doc":
        # legacy .doc (binary) is not reliably parsed with python-docx.
        # We intentionally do not attempt to parse .doc here.
        # Ask user to convert to .docx or upload .txt instead.
        return "", "doc"
    else:
        return "", "unknown"

# ---------- Public landing & auth routes ----------
@app.route('/', methods=['GET'])
def landing():
    """Landing page - public route."""
    try:
        return render_template('landing.html')
    except Exception as e:
        print(f"Landing page error: {e}")
        return render_template('error.html', error="Unable to load landing page", code=500), 500

@app.route('/health', methods=['GET'])
def health():
    """Health check endpoint for monitoring."""
    try:
        # Try to get a database connection to verify it's working
        conn = get_conn()
        conn.close()
        return {'status': 'healthy', 'database': 'ok'}, 200
    except Exception as e:
        print(f"Health check failed: {e}")
        return {'status': 'unhealthy', 'database': 'error', 'error': str(e)}, 503

@app.route('/register', methods=['GET','POST'])
def register():
    if request.method == 'POST':
        try:
            full_name = request.form.get('full_name','').strip()
            email = request.form.get('email','').strip().lower()
            password = request.form.get('password','')
            confirm = request.form.get('confirm','')
            role = request.form.get('role','student')
            college = request.form.get('college','').strip() or None
            department = request.form.get('department','').strip() or None
            year = request.form.get('year','').strip() or None
            enrollment_no = request.form.get('enrollment_no','').strip() or None
            phone = request.form.get('phone','').strip() or None

            if not full_name or not email or not password:
                flash('Full name, email and password are required.', 'error')
                return redirect(url_for('register'))
            if password != confirm:
                flash('Passwords do not match.', 'error')
                return redirect(url_for('register'))

            if get_user_by_email(email):
                flash('An account with this email already exists.', 'error')
                return redirect(url_for('register'))

            pw_hash = generate_password_hash(password)
            user_id = create_user(full_name, email, pw_hash, role, college, department, year, enrollment_no, phone, None)
            session['user_id'] = user_id
            flash('Registration successful. You are now logged in.', 'success')
            return redirect(url_for('index'))
        except Exception as e:
            print("DB register error:", e)
            flash('Registration failed due to server error.', 'error')
            return redirect(url_for('register'))

    return render_template('register.html')

@app.route('/login', methods=['GET','POST'])
def login():
    if request.method == 'POST':
        try:
            email = request.form.get('email','').strip().lower()
            password = request.form.get('password','')
            print(f"DEBUG: Login attempt for email: {email}")
            user = get_user_by_email(email)
            print(f"DEBUG: User found: {user is not None}")
            if user is None or not check_password_hash(user['password_hash'], password):
                print(f"DEBUG: Password check failed for user: {user}")
                flash('Incorrect email or password.', 'error')
                return redirect(url_for('login'))
            session.clear()
            session['user_id'] = user['id']
            print(f"DEBUG: Login successful for user: {user['id']}")
            flash('Logged in successfully.', 'success')
            next_page = request.args.get('next') or url_for('index')
            return redirect(next_page)
        except Exception as e:
            print(f"Login error: {e}")
            flash('Login failed due to server error.', 'error')
            return redirect(url_for('login'))
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    flash('Logged out.', 'success')
    return redirect(url_for('landing'))

from db import get_all_users

@app.route('/admin/dashboard')
@login_required
@login_required
def admin_dashboard():
    users = get_all_users()
    return render_template('admin_dashboard.html', users=users)


def admin_required(view):
    @functools.wraps(view)
    def wrapped_view(*args, **kwargs):
        if g.user is None or g.user.get('role') != 'admin':
            flash("Admin access required.", "error")
            return redirect(url_for('login'))
        return view(*args, **kwargs)
    return wrapped_view

@app.route('/admin', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        username = request.form.get('username','')
        password = request.form.get('password','')

        user = get_user_by_email(username)

        if user and user['role'] == 'admin' and check_password_hash(user['password_hash'], password):
            session.clear()
            session['user_id'] = user['id']
            return redirect(url_for('admin_dashboard'))

        flash("Invalid admin credentials.", "error")
        return redirect(url_for('admin_login'))

    return render_template('admin_login.html')


# ---------- Dashboard and syllabus processing ----------
@app.route("/index", methods=["GET"])
@login_required
def index():
    return render_template("index.html")

@app.route("/process", methods=["POST"])
@login_required
def process():
    """
    Accept either:
    - file upload (.txt or .docx) via 'syllabus_file'
    - or plain text in 'syllabus_text' textarea
    Priority: file upload (if provided and valid) -> textarea -> error
    """
    try:
        title = request.form.get("title", "My Syllabus")
        syllabus_text = (request.form.get("syllabus_text") or "").strip()

        # File upload handling
        file = request.files.get("syllabus_file")
        uploaded_text = ""
        upload_type = "none"
        if file and file.filename:
            if not allowed_file(file.filename):
                flash("Unsupported file type. Please upload .txt or .docx (or convert .doc to .docx).", "error")
                return redirect(url_for("index"))
            try:
                uploaded_text, upload_type = extract_text_from_upload(file)
            except ValueError as ve:
                flash(str(ve), "error")
                return redirect(url_for("index"))
            except Exception as e:
                print("Upload parsing error:", e)
                flash("Failed to parse uploaded file. Try .txt or .docx formats.", "error")
                return redirect(url_for("index"))

            if upload_type == "doc":
                flash("Legacy .doc files are not supported. Please save as .docx or .txt and upload again.", "error")
                return redirect(url_for("index"))

            if not uploaded_text:
                flash("Uploaded file seemed empty or unparsable. Please check the file and try again.", "error")
                return redirect(url_for("index"))

        # Final syllabus source selection
        final_text = uploaded_text if uploaded_text else syllabus_text
        if not final_text:
            flash("Please paste syllabus text or upload a .txt / .docx file.", "error")
            return redirect(url_for("index"))

        # Persist syllabus
        sid = insert_syllabus(title, final_text)

        # NLP: extract units->topics (function already robust)
        units = extract_topics_per_unit(final_text, topics_per_unit=6)
        if not units:
            flash("Could not extract any units/topics from your syllabus. Please add more details and try again.", "error")
            return redirect(url_for("index"))

        # For each unit/topic -> query youtube and store
        for u in units:
            try:
                unit_id = insert_unit(sid, u["unit_no"], u.get("unit_title"))
                if not u.get("topics"):
                    u["topics"] = [{"text": "overview", "weight": 1.0}]
                for t in u["topics"]:
                    topic_id = insert_topic(unit_id, t["text"], t["weight"])
                    query = f"{t['text']} lecture tutorial"
                    try:
                        results = search_and_rank(Config.YT_API_KEY, query, max_results=12)
                        for r in results[:12]:
                            insert_video(topic_id, r)
                    except Exception as e:
                        # logged but don't break
                        print("YouTube API error (ignored):", e)
            except Exception as e:
                print(f"Error processing unit {u.get('unit_no')}: {e}")
                continue

        return redirect(url_for("results", sid=sid))
    except Exception as e:
        print(f"Process error: {e}")
        flash("Processing failed due to server error. Please try again.", "error")
        return redirect(url_for("index"))

@app.route("/results/<int:sid>", methods=["GET"])
@login_required
def results(sid):
    try:
        duration = request.args.get("duration", "all")
        difficulty = request.args.get("difficulty", "all")
        min_rating = request.args.get("min_rating", None)
        filters = {
            "duration": duration if duration in ("short","medium","long") else None,
            "difficulty": difficulty if difficulty in ("beginner","intermediate","advanced") else None,
            "min_rating": float(min_rating) if min_rating else None
        }
        units = fetch_units_topics_videos(sid, filters)
        return render_template("results.html", sid=sid, units=units, filters=filters)
    except Exception as e:
        print(f"Results error: {e}")
        flash("Failed to load results. Please try again.", "error")
        return redirect(url_for("index"))

# For local development only - not used in Vercel deployment
if __name__ == "__main__":
    app.run(debug=True)
